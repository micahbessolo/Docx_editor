from datetime import date
import datetime
from docx import Document
document = Document('C:/Users/besso/VSC-Python/nonspouse.docx')
paragraphs = document.paragraphs


# Add Today's Date to paragraph
today = date.today()
datee = datetime.datetime.strptime(str(today), "%Y-%m-%d")
# turns month number into full month name
datetime_object = datetime.datetime.strptime(str(datee.month), "%m")
full_month_name = datetime_object.strftime("%B")
fulldate = str(datee.day) + " " + full_month_name + " " + str(datee.year)
document.paragraphs[2].add_run(fulldate)

# Beneficiary Name
beneficiary_name = input('beneficiaries name: ')
# Add Beneficiary's name to paragraph
document.paragraphs[4].add_run(beneficiary_name)

# Add Street Address to paragraph
street_address = input("beneficiary's street address: ")
document.paragraphs[5].add_run(street_address)

# Add City, State, Zip Code to paragraph
city_state_zip = input("city state zip: ")
document.paragraphs[6].add_run(city_state_zip)

# Add decedent's name to paragraph
decedents_name = input("decedent's name: ")
document.paragraphs[8].runs[2].add_text(decedents_name + "'s")

# Add account number to paragraph
account_number = input("account number: ")
# replace middle numbers with Xs
number_list = list(account_number)
number_list[4:9] = 'XXXXX'
new_number = ""
document.paragraphs[8].runs[6].add_text(new_number.join(number_list))

# Add Mr./Ms. Bene Last Name to paragraph
bene_title = input("beneficiary title: ")
document.paragraphs[12].runs[1].add_text(decedents_name + "'s")

# Add representative's name to paragraph
rep_name = input("rep's name: ")
document.paragraphs[14].runs[3].add_text(rep_name)

# Add representative's number to paragraph
rep_number = input("rep's phone number: ")
document.paragraphs[14].runs[5].add_text(rep_number)

# Add Service Agent's Name to paragraph
sc_agent = 'Micah Bessolo'
document.paragraphs[-4].add_run(sc_agent)

# saves updates to file
document.save("C:/Users/besso/VSC-Python/" + beneficiary_name + 'delete.docx')
