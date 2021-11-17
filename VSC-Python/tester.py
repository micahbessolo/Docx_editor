from datetime import date
import datetime
from docx import Document
document = Document('C:/Users/besso/VSC-Python/nonspouse.docx')
paragraphs = document.paragraphs


# Add Today's Date to paragraph
today = date.today()
datee = datetime.datetime.strptime(str(today), "%Y-%m-%d")
# turns month number into full month name
datetime_object = datetime.datetime.strptime(str(datee.month), "%m")
full_month_name = datetime_object.strftime("%B")
fulldate = str(datee.day) + " " + full_month_name + " " + str(datee.year)
document.paragraphs[2].add_run(fulldate)

# Beneficiary Name
beneficiary_name = 'Micah Bessolo'
# Add Beneficiary's name to paragraph
document.paragraphs[4].add_run(beneficiary_name)

# Add Street Address to paragraph
street_address = "2196 Edgebrook Pl"
document.paragraphs[5].add_run(street_address)

# Add City, State, Zip Code to paragraph
city_state_zip = 'Hayward, CA 94541'
document.paragraphs[6].add_run(beneficiary_name)

# Add decedent's name to paragraph
decedents_name = 'Dead Person Name'
document.paragraphs[8].runs[2].add_text(decedents_name + "'s")

# Add account number to paragraph
account_number = "510512047-01"
# replace middle numbers with Xs
number_list = list(account_number)
number_list[4:9] = 'XXXXX'
new_number = ""
document.paragraphs[8].runs[3].add_text(new_number.join(number_list))

# Add Mr./Ms. Bene Last Name to paragraph
bene_title = "Mr. Bessolo"
document.paragraphs[12].runs[1].add_text(decedents_name + "'s")

# Add representative's name to paragraph
rep_name = 'Barabara Seaman'
document.paragraphs[14].runs[1].add_text(rep_name)

# Add representative's number to paragraph
rep_number = '(510) 512-0476'
document.paragraphs[14].runs[4].add_text(rep_number)

# Add Service Agent's Name to paragraph
sc_agent = 'Micah Bessolo'
document.paragraphs[-4].add_run(sc_agent)

# saves updates to file
document.save(beneficiary_name + 'delete.docx')
